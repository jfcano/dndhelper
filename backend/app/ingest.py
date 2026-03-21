from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import hashlib
import json
import logging
from typing import Callable

from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from tqdm import tqdm

from backend.app.config import get_settings
from backend.app.vector_store import get_vector_store

logger = logging.getLogger(__name__)


def _load_langchain_documents(path: Path) -> list[Document]:
    """Carga páginas/bloques según extensión (.pdf, .txt, .docx)."""
    suf = path.suffix.lower()
    if suf == ".pdf":
        return PyPDFLoader(str(path)).load()
    if suf == ".txt":
        raw = path.read_bytes()
        text = raw.decode("utf-8-sig", errors="replace")
        text = _sanitize_utf8(text.strip())
        if not text:
            return []
        return [Document(page_content=text, metadata={"source": str(path.resolve()), "page": 1})]
    if suf == ".docx":
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        parts: list[str] = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                parts.append(_sanitize_utf8(t))
        for table in doc.tables:
            for row in table.rows:
                cells = [_sanitize_utf8(c.text.strip()) for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n\n".join(parts)
        text = text.strip()
        if not text:
            return []
        return [Document(page_content=text, metadata={"source": str(path.resolve()), "page": 1})]
    raise ValueError(f"Formato no soportado: {suf}")


def _vectors_still_present_for_pdf(vs, pdf: Path) -> bool:
    """
    Comprueba si hay chunks en PGVector para este PDF (misma colección).
    Evita omitir la ingesta cuando el manifiesto local dice «duplicado» pero la BD
    se vació o se recreó (p. ej. migraciones) y el índice ya no existe.
    """
    src_resolved = str(pdf.resolve())
    alts = {src_resolved, str(pdf)}
    try:
        found = vs.similarity_search(" ", k=1, filter={"source": src_resolved})
        if found:
            return True
    except Exception:
        logger.debug("similarity_search con filter falló; probando sin filter.", exc_info=True)
    try:
        for d in vs.similarity_search(" ", k=48):
            meta = d.metadata or {}
            s = meta.get("source")
            if s in alts:
                return True
    except Exception:
        logger.warning("No se pudo comprobar embeddings en PGVector; se reindexará.", exc_info=True)
        return False
    return False

# Tamaño de lote para indexar (cada lote implica llamadas a la API de embeddings de OpenAI)
_INGEST_BATCH_SIZE = 32


class IngestCancelledError(Exception):
    """El usuario canceló el trabajo (estado «cancelled» en BD); el worker debe limpiar y no finalizar como éxito."""


@dataclass(frozen=True)
class IngestResult:
    file_path: str
    pdf_sha256: str
    chunks_indexed: int
    collection: str
    manifest_path: str
    skipped_duplicate: bool = False


def _sanitize_utf8(s: str) -> str:
    """Normaliza texto para almacenamiento/embeddings.

    - Reemplaza surrogates y otros caracteres no encodables en UTF-8.
    - Elimina bytes NUL ('\\x00'), que Postgres no permite en campos TEXT.
    """
    s = s.encode("utf-8", errors="replace").decode("utf-8")
    return s.replace("\x00", "")


def _sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for block in iter(lambda: f.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def _manifest_path(project_root: Path) -> Path:
    # Archivo local para evitar re-ingestas innecesarias. No está relacionado con el vector store.
    storage_dir = project_root / "backend" / "storage"
    storage_dir.mkdir(parents=True, exist_ok=True)
    return storage_dir / "ingest_manifest.json"


def _load_manifest(project_root: Path) -> dict:
    p = _manifest_path(project_root)
    if not p.exists():
        return {}
    return json.loads(p.read_text(encoding="utf-8"))


def _save_manifest(project_root: Path, manifest: dict) -> None:
    p = _manifest_path(project_root)
    p.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")


def remove_manifest_entries_for_collection(*, collection_name: str) -> int:
    """Elimina entradas del manifiesto de ingesta CLI cuya clave empieza por ``collection_name:``."""
    settings = get_settings()
    manifest = _load_manifest(settings.project_root)
    prefix = f"{collection_name}:"
    to_del = [k for k in list(manifest.keys()) if isinstance(k, str) and k.startswith(prefix)]
    for k in to_del:
        del manifest[k]
    if to_del:
        _save_manifest(settings.project_root, manifest)
    return len(to_del)


def ingest_document(
    document_path: str,
    *,
    collection_name: str,
    force: bool = False,
    show_progress: bool = True,
    progress_callback: Callable[[str, int, int], None] | None = None,  # phase: load|split|index|unchanged
    cancel_check: Callable[[], bool] | None = None,
) -> IngestResult:
    """
    Indexa un documento (PDF, TXT o DOCX) en Postgres (pgvector) con LangChain.
    Los embeddings son siempre de **OpenAI** (`OpenAIEmbeddings`, modelo `OPENAI_EMBEDDINGS_MODEL` vía `get_embeddings`).
    ``collection_name``: colección PGVector (p. ej. una por usuario vía ``rag_collection_name_for_owner``).
    show_progress: si True, usa tqdm en consola para la fase de indexación.
    progress_callback(phase, current, total): opcional; phase in ("load", "split", "index", "unchanged").
    cancel_check: si devuelve True, se interrumpe la ingesta (levanta IngestCancelledError) entre fases/lotes.
    """
    settings = get_settings()

    def _abort_if_cancelled() -> None:
        if cancel_check and cancel_check():
            raise IngestCancelledError()

    pdf = Path(document_path)
    if not pdf.exists():
        raise FileNotFoundError(f"No existe el fichero: {pdf}")
    _abort_if_cancelled()

    pdf_sha = _sha256_file(pdf)
    manifest = _load_manifest(settings.project_root)
    key = f"{collection_name}:{str(pdf.resolve())}"

    if not force and manifest.get(key, {}).get("pdf_sha256") == pdf_sha:
        vs_check = get_vector_store(collection_name=collection_name)
        if _vectors_still_present_for_pdf(vs_check, pdf):
            if progress_callback:
                progress_callback("unchanged", 1, 1)
            mp = _manifest_path(settings.project_root)
            return IngestResult(
                file_path=str(pdf),
                pdf_sha256=pdf_sha,
                chunks_indexed=0,
                collection=collection_name,
                manifest_path=str(mp),
                skipped_duplicate=True,
            )
        logger.info(
            "Manifiesto marca duplicado pero no hay vectores en «%s» para %s; reindexando.",
            collection_name,
            pdf.name,
        )

    _abort_if_cancelled()

    def _report(phase: str, current: int, total: int) -> None:
        if progress_callback:
            progress_callback(phase, current, total)

    # Fase 1: cargar documento
    if show_progress:
        tqdm.write(f"Cargando documento: {pdf.name}")
    _abort_if_cancelled()
    try:
        docs = _load_langchain_documents(pdf)
    except ValueError as e:
        raise RuntimeError(str(e)) from e
    except Exception as e:
        raise RuntimeError(f"No se pudo leer el documento: {e}") from e
    _report("load", len(docs) if docs else 1, max(len(docs), 1))
    if show_progress:
        tqdm.write(f"  Bloques cargados: {len(docs)}")

    # Fase 2: dividir en chunks
    if show_progress:
        tqdm.write("Dividiendo en chunks...")
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
    )
    chunks = splitter.split_documents(docs)
    # Excluir chunks vacíos y normalizar page_content (UTF-8 válido, sin surrogates)
    valid = []
    for d in chunks:
        raw = getattr(d, "page_content", None)
        if raw is None:
            continue
        text = str(raw).strip()
        if not text:
            continue
        text = _sanitize_utf8(text)
        d.page_content = text
        valid.append(d)
    chunks = valid
    _report("split", len(chunks), len(chunks))
    if show_progress:
        tqdm.write(f"  Chunks generados: {len(chunks)}")

    # Metadatos mínimos útiles para citar origen (ruta canónica = misma clave que el manifiesto)
    src_tag = str(pdf.resolve())
    for d in chunks:
        d.metadata = d.metadata or {}
        d.metadata.setdefault("source", src_tag)

    _abort_if_cancelled()

    vs = get_vector_store(collection_name=collection_name)
    # Si re-ingestamos, limpiamos la colección para no duplicar
    if force or manifest.get(key):
        try:
            vs.delete_collection()
        except Exception:
            pass
        vs = get_vector_store(collection_name=collection_name)

    # Fase 3: indexar en lotes (embeddings vía OpenAIEmbeddings → API OpenAI, ver get_embeddings)
    if not chunks:
        _report("index", 0, 0)
    if chunks:
        logger.info(
            "Ingesta RAG: embeddings OpenAI (modelo=%s, colección=%s, chunks=%d)",
            settings.openai_embeddings_model,
            collection_name,
            len(chunks),
        )
    if show_progress:
        tqdm.write("Indexando embeddings (OpenAI API; puede tardar según tamaño y red)...")
    batch_size = _INGEST_BATCH_SIZE
    for start in tqdm(
        range(0, len(chunks), batch_size),
        desc="Indexando chunks",
        total=(len(chunks) + batch_size - 1) // batch_size,
        unit="lote",
        disable=not show_progress,
    ):
        _abort_if_cancelled()
        batch = chunks[start : start + batch_size]
        vs.add_documents(batch)
        _report("index", min(start + len(batch), len(chunks)), len(chunks))

    manifest[key] = {
        "pdf_sha256": pdf_sha,
        "pdf_path": str(pdf.resolve()),
        "collection": collection_name,
        "chunks": len(chunks),
    }
    _save_manifest(settings.project_root, manifest)
    mp = _manifest_path(settings.project_root)

    return IngestResult(
        file_path=str(pdf),
        pdf_sha256=pdf_sha,
        chunks_indexed=len(chunks),
        collection=collection_name,
        manifest_path=str(mp),
        skipped_duplicate=False,
    )


# Compatibilidad con scripts que aún nombran la función «ingest_pdf».
ingest_pdf = ingest_document

